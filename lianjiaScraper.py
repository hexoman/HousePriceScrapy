# -*- coding: UTF-8 -*-
import sys
reload(sys)
sys.setdefaultencoding('utf8')
import requests
import urllib2
from lxml import html
import matplotlib.pyplot as plt
import pyecharts
from pyecharts import Pie
from pyecharts import Bar
import numpy as np
from docx import Document
from docx.shared import Inches
from PIL import Image

CITY = ''
DATACOUNT = 0
TotalPage = 330
fiveDownLevel = '5k以下'
fiveKLevel = '5k-1w'
OneOneLevel = '1w-1.5w'
OneTwoLevel = '1.5w-2w'
TwoOneLevel = '2w-2.5w'
TwoTwoLevel = '2.5w-3w'
ThreeAboveLevel = '3w以上'
addressDict = {}
priceDict = {}
# TwoOneLevel = '3w以下'
# threeHalfLevel = '3w-3.5w'
# threeToFourLevel = '3.5w-4w'
# halfFourLevel = '4w-4.5w'
# fourToFiveLevel = '4.5w-5w'
# fiveAboveLevel = '5w以上'

def crow(i):
    print('正在爬取第{}页,还剩{}页'.format(i,TotalPage-i))
    # 拼装第i页的网址
    url = 'https://sh.lianjia.com/ershoufang/pg%s/' % str(i)

    # 构造模拟请求的客户端信息
    headers = {'User-Agent':'Mozilla/5.0 (Windows NT 6.2; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/62.0.3202.62 Safari/537.36'}
    req = urllib2.Request(url=url,headers=headers)#这里要注意，必须使用url=url，headers=headers的格式，否则回报错，无法连接字符
    response = urllib2.urlopen(req)#注意，这里要用req，不然就被添加userag
    rst = response.read().decode("utf-8")
    rst=html.fromstring(rst)
    path = '/html/body/div[4]/div[1]/ul/li'
    parseHtml(path, rst)
    global CITY
    st = rst.xpath('/html/body/div[2]/div[1]/div[1]/ul/li[1]/a/@title')[0].strip()
    CITY = st[0:st.find(u'在售二手房')]

#解析网址
def parseHtml(path, rst):
    datas = rst.xpath(path)
    global DATACOUNT 
    DATACOUNT += len(datas)
    for data in datas:
        title=data.xpath('/html/body/div[4]/div[1]/ul/li[1]/div[1]/div[1]/a/text()')
        area=data.xpath('/html/body/div[4]/div[1]/ul/li[1]/div[1]/div[3]/div/a/text()')
        houseInfo=data.xpath('/html/body/div[4]/div[1]/ul/li[1]/div[1]/div[2]/div/text()')
        flood=data.xpath('/html/body/div[4]/div[1]/ul/li[1]/div[1]/div[3]/div/text()')
        followInfo=data.xpath('/html/body/div[4]/div[1]/ul/li[1]/div[1]/div[4]/text()')
        totalPrice=data.xpath('/html/body/div[4]/div[1]/ul/li[1]/div[1]/div[6]/div[1]/span/text()')
        unitPrice=data.xpath('/html/body/div[4]/div[1]/ul/li[1]/div[1]/div[6]/div[2]/span/text()')
        # address=data.xpath('div/div[1]/p/text()')
        # saveTolocal(company,companyDesc,title,address,price)
        bulidDatas(area,totalPrice,unitPrice)


#组装分析的数据
def bulidDatas(area,totalPrice,unitPrice): 
    #区域
    ar = area[0].strip()
    if ar in addressDict:
       addressDict[ar].append(1)
    else:
       addressDict[ar] = [1] 
    
    #均价
    p = unitPrice[0].strip()
    sp = p.find('单价')
    kp = p.find('元')
    startP = int(p[2:kp])   
    # if startP < 30000:
    #     p = threeLevel
    # elif startP >= 30000 and startP < 35000:
    #     p = threeHalfLevel
    # elif startP >= 35000 and startP < 40000:
    #     p = halfFourLevel
    # elif startP >= 40000 and startP < 45000:
    #     p = halfFourLevel      
    # elif startP >= 45000 and startP < 50000:
    #     p = fourToFiveLevel
    # elif startP >= 50000:
    #     p = fiveAboveLevel  
    if startP < 5000:
        p = fiveDownLevel
    elif startP >= 5000 and startP < 10000:
        p = fiveKLevel
    elif startP >= 10000 and startP < 15000:
        p = OneOneLevel
    elif startP >= 15000 and startP < 20000:
        p = OneTwoLevel      
    elif startP >= 20000 and startP < 25000:
        p = TwoOneLevel
    elif startP >= 25000 and startP < 30000:
        p = TwoTwoLevel
    elif startP >= 30000:
        p = ThreeAboveLevel 
    if p in priceDict:
       priceDict[p].append(1)
    else:
       priceDict[p] = [1] 

#绘制地址分布图表
def drawAddressSheet():
    labels = []
    sizes = []
    for key in addressDict:
        labels.append(key)
        vList = addressDict[key]
        sizes.append(len(vList))
    drawBarGraph(sizes,labels,'%s区域范围分布图'%CITY,'机会数量','区域范围分布图')


#绘制薪资分布图表
def drawPriceLevelSheet():
    labels = []
    sizes = []
    for key in priceDict:
        labels.append(key)
        vList = priceDict[key]
        sizes.append(len(vList))
    drawBarGraph(sizes,labels,'%s均价分布饼状图'%CITY,'薪资范围','均价分布柱形图')    

def drawBarGraph(sizes,labels,title,category,imgName):
    #饼状图
    pie = Pie(title,title_pos='center')
    # pie.add(category,labels,sizes,is_lable_show=True)
    pie.add(category, labels, sizes, radius=[40, 75],
        label_text_color=None,  #标签字体的颜色
        is_label_show=True,
        legend_orient='vertical',
        legend_pos='left')
    pie.render(path='./%s.png' % imgName)

    #柱形图
    # bar = Bar(title,'Theme',width=1100)
    # bar.use_theme('dark')
    # bar.add(category,labels,sizes)
    # bar.render(path='./%s.png' % imgName)


# def drawCanvs(sizes,labels,colors):
#     print(sizes)
#     plt.figure(figsize=(9,9)) #调节图形大小
#     explode = (0.01,0,0,0) #将某一块分割出来，值越大分割出的间隙越大
#     plt.axes(aspect=1)  # 设置X轴 Y轴比例 pctdistance=0.8,
#     # plt.legend(loc=7, bbox_to_anchor=(1.2, 0.80), ncol=3, fancybox=True, shadow=True, fontsize=5)
#     patches,text1,text2 = plt.pie(sizes, labels=labels,colors=colors, autopct='%3.1f %%',
#         shadow=False, labeldistance=1.1, startangle=0)
#     plt.axis('equal')
#     plt.show()


#保存到本地文本
def saveTolocal(company,companyDesc,title,address,price): 
    with open('boss.txt','a')as f:
        f.write(company[0]+'\n')
        f.write(companyDesc[0]+'\n')
        f.write(title[0]+'\n')
        f.write(address[0]+'\n')
        f.write(price[0]+'\n')
        f.write('\n'*3)

def saveToWord():
    doc = Document()
    doc.add_paragraph(u'%s二手房行业分析报告'%CITY + '\n'*2) 
    doc.add_picture(u'%s区域范围分布图.png'%CITY, width=Inches(8))
    doc.add_paragraph('\n'*2) 
    doc.add_picture(u'%s均价分布饼状图.png'%CITY, width=Inches(8))
    doc.add_paragraph('\n'*2) 
    doc.save(u'行业分析报告.docx')


#开始抓取数据
def startScraper():
   for i in range(TotalPage):
       print('爬取数据中...')
       crow(i)
   print('绘制中...')
   drawAddressSheet()
   drawPriceLevelSheet()
   print('绘制完成')
   print('总共爬取{}{}条数据'.format(CITY,str(DATACOUNT)))
   print('生成报告文档...')
   saveToWord()
   print('生成成功')
startScraper()
